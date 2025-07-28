import os
import subprocess
from datetime import datetime
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches
from docxtpl import DocxTemplate

from app.controller.properties_controller import SingleCmaController
from app.routing.services import PropertyService, ReportService
from app.utils.mapbox_utils import gen_static_map_url, get_mapbox_static_map
from config import EVIDENCE_DIR

try:
    from comtypes import client
except ImportError:
    client = None


class EvidenceDocument:

    TEMPORARY_TEMPLATE_NAME = 'temporary_doc_template.docx'
    TEMPORARY_MAPBOX_IMAGE_NAME = 'temporary_mapbox_image.png'

    def __init__(self, template_path):
        self._docx = DocxTemplate(template_path)

        # make sure evidence directory created
        base_dir = EVIDENCE_DIR
        base_dir.mkdir(exist_ok=True, parents=True)

    def gen_petition_document(self, case):

        # get case object
        case_client = case.client
        subject = case.property
        apn = subject.apn

        # get the property assessment
        assessment = PropertyService.get_property_assessment(subject.id, tax_year=case.tax_year)

        # create CMA controller
        controller = SingleCmaController(subject=subject, subject_assessment=assessment, mass_cma=False)
        controller.compute_cma()
        cma_results = controller.get_derived_assessment_results()
        subject = controller.subject_property

        context = {
            'county': subject.county,
            'address': subject.address,
            'year_built': subject.year,
            'beds': subject.bedrooms,
            'baths': subject.full_baths,
            'gla_sqft': subject.gla_sqft,
            'apn': apn,
            'full_name': case_client.full_name,
            'assessment_value': ReportService.format_price(cma_results.get('current_assessment_value')),
            'proposed_value': ReportService.format_price(cma_results.get('proposed_assessment_value')),
            'date_submitted': str(datetime.now().strftime('%m/%d/%y'))
        }

        self._docx.render(context)
        (EVIDENCE_DIR / apn).mkdir(exist_ok=True, parents=True)
        doc_file_path = EVIDENCE_DIR / apn / f"pet_evid_{apn}.docx"
        self._docx.save(doc_file_path)

        self.doc2pdf(doc_file_path)
        os.remove(doc_file_path)

        pdf_file_path = f'{Path(doc_file_path).as_posix()[:-5]}.pdf'
        return pdf_file_path

    @staticmethod
    def get_mapbox_image_path():
        return EVIDENCE_DIR / EvidenceDocument.TEMPORARY_MAPBOX_IMAGE_NAME

    @staticmethod
    def get_document_path():
        return EVIDENCE_DIR / EvidenceDocument.TEMPORARY_TEMPLATE_NAME

    def add_static_map(self, context):

        from manage import app
        with app.app_context():
            url = gen_static_map_url(context.get('pins'))
            _ = get_mapbox_static_map(url, output_path=self.get_mapbox_image_path()) # noqa

        paragraph = self._docx.add_paragraph()

        r = paragraph.add_run()
        r.add_break(WD_BREAK.LINE)
        # height = Inches(3.7)
        r.add_picture(self.get_mapbox_image_path().as_posix(), width=Inches(6.5))
        last_paragraph = self._docx.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_print_scan(self, apn):
        # add page from 'broward' website
        paragraph = self._docx.add_paragraph()
        r = paragraph.add_run()
        r.add_break(WD_BREAK.PAGE)

        import tempfile
        from pdf2image import convert_from_path

        with tempfile.TemporaryDirectory(dir=EVIDENCE_DIR) as path:
            paths = convert_from_path((EVIDENCE_DIR / apn / f'{apn}.pdf').as_posix(), fmt='png',
                                      paths_only=True, output_folder=path)

            r.add_picture(paths[0])
            last_paragraph = self._docx.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def doc2pdf_linux(self, docx):
        """
        Convert a doc/docx document to pdf format (linux only, requires libreoffice)
        :param docx: The path to document
        """
        output_dir = Path(docx).parent.absolute()
        try:
            subprocess.run(
                ['libreoffice', '--convert-to', 'pdf', str(docx), '--outdir', output_dir.as_posix()],
                stdout=subprocess.PIPE,
                universal_newlines=True
            )
        except Exception as e:
            print(e)

    def doc2pdf(self, docx=None):
        """
        Convert a doc/docx document to pdf format
        :param docx: The path to document
        """
        if docx is None:
            docx = self.get_document_path().as_posix()

        # bug fix - searching files in windows/system32
        docx = os.path.abspath(docx)
        if client is None:
            return self.doc2pdf_linux(docx)
        name, ext = os.path.splitext(docx)

        word = client.CreateObject('Word.Application')
        word_doc = word.Documents.Open(docx)
        try:
            word_doc.SaveAs(name + '.pdf', FileFormat=17)
            return name + '.pdf'
        except Exception:
            raise
        finally:
            word_doc.Close()
            word.Quit()
